import io
import re
from PyPDF2 import PdfReader
from fastapi import UploadFile
import docx
from PIL import Image
import pytesseract
from sklearn.feature_extraction.text import CountVectorizer
from collections import Counter
import logging

logger = logging.getLogger(__name__)

# Set the path to the Tesseract executable
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

async def get_text_and_metadata(file: UploadFile, doc_type: str):
    content = await file.read()
    text = ""
    metadata = {}
    images = []

    if doc_type == "pdf":
        try:
            pdf_reader = PdfReader(io.BytesIO(content))
            metadata = pdf_reader.metadata
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
                if '/XObject' in page['/Resources']:
                    xObject = page['/Resources']['/XObject'].get_object()
                    for obj in xObject:
                        if xObject[obj]['/Subtype'] == '/Image':
                            try:
                                img_data = xObject[obj]._data
                                images.append(Image.open(io.BytesIO(img_data)))
                            except:
                                logger.warning("Skipped unidentified image in PDF")
        except Exception as e:
            logger.error(f"Error processing PDF: {str(e)}")
            raise ValueError(f"Error processing PDF {file.filename}: {str(e)}")

    elif doc_type == "docx":
        try:
            doc = docx.Document(io.BytesIO(content))
            for para in doc.paragraphs:
                text += para.text + "\n"
            metadata = {"docx_length": len(doc.paragraphs)}
        except Exception as e:
            logger.error(f"Error processing DOCX: {str(e)}")
            raise ValueError(f"Error processing DOCX {file.filename}: {str(e)}")

    metadata.update({
        "experience": extract_experience(text),
        "location": extract_location(text),
        "job_category": extract_job_category(text)
    })

    return text, metadata, images

def extract_experience(text):
    experience_pattern = r'\b(\d+)\s*(?:years?|yrs?)\s*(?:of\s*)?experience\b'
    matches = re.findall(experience_pattern, text, re.IGNORECASE)
    return int(matches[0]) if matches else 0

def extract_location(text):
    locations = ["new york", "san francisco", "london", "berlin", "tokyo"]
    for location in locations:
        if location.lower() in text.lower():
            return location
    return "Unknown"

def extract_job_category(text):
    vectorizer = CountVectorizer(stop_words='english', max_features=50)
    vector = vectorizer.fit_transform([text])
    words = vectorizer.get_feature_names_out()
    word_counts = vector.toarray()[0]
    word_freq = Counter(dict(zip(words, word_counts)))
    
    category_keywords = {
        "INFORMATION-TECHNOLOGY": ["it", "software", "developer", "programming", "web", "database", "network", "security", "cloud", "devops"],
        "BUSINESS-DEVELOPMENT": ["business", "development", "strategy", "growth", "partnerships", "sales", "marketing"],
        "ADVOCATE": ["lawyer", "legal", "attorney", "law", "litigation", "counsel"],
        "CHEF": ["chef", "cook", "culinary", "kitchen", "food", "restaurant"],
        "ENGINEERING": ["engineer", "mechanical", "electrical", "civil", "chemical", "software"],
        "ACCOUNTANT": ["accountant", "accounting", "auditor", "bookkeeper", "financial"],
        "FINANCE": ["finance", "investment", "banking", "analyst", "portfolio", "trading"],
        "FITNESS": ["fitness", "trainer", "gym", "exercise", "health", "wellness"],
        "AVIATION": ["pilot", "aviation", "aircraft", "flight", "airline"],
        "SALES": ["sales", "account", "representative", "business", "client"],
        "BANKING": ["banking", "bank", "credit", "loan", "financial", "teller"],
        "HEALTHCARE": ["healthcare", "medical", "doctor", "nurse", "patient", "hospital"],
        "CONSULTANT": ["consultant", "consulting", "advisor", "strategy", "management"],
        "CONSTRUCTION": ["construction", "builder", "contractor", "project", "site"],
        "PUBLIC-RELATIONS": ["pr", "public", "relations", "communications", "media"],
        "HR": ["hr", "human", "resources", "recruitment", "talent", "personnel"],
        "DESIGNER": ["designer", "design", "graphic", "ux", "ui", "creative"],
        "ARTS": ["artist", "creative", "painter", "sculptor", "performer"],
        "TEACHER": ["teacher", "educator", "instructor", "professor", "tutor"],
        "APPAREL": ["fashion", "clothing", "textile", "garment", "designer"],
        "DIGITAL-MEDIA": ["digital", "media", "content", "social", "marketing"],
        "AGRICULTURE": ["agriculture", "farming", "crop", "livestock", "agronomy"],
        "AUTOMOBILE": ["automobile", "automotive", "mechanic", "car", "vehicle"],
        "BPO": ["bpo", "outsourcing", "call", "center", "customer", "service"]
    }
    
    category_scores = {category: sum(word_freq[word] for word in keywords if word in word_freq) 
                       for category, keywords in category_keywords.items()}
    
    if category_scores:
        best_category = max(category_scores, key=category_scores.get)
        if category_scores[best_category] > 0:
            return best_category
    
    return "OTHER"

def extract_highlights(resume_content: str, query: str) -> str:
    query_keywords = set(query.lower().split())
    resume_words = resume_content.lower().split()
    highlights = [word for word in resume_words if word in query_keywords]
    return ", ".join(highlights[:5])